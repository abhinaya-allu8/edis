"""
EDIS — DOCX Parser
Extracts text from Word documents. Preserves paragraph and table content.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import List

from docx import Document


@dataclass
class ParsedDocument:
    source: str
    doc_type: str
    pages: List[dict] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p["text"] for p in self.pages if p["text"].strip())

    @property
    def metadata(self) -> dict:
        return {"source": self.source, "doc_type": self.doc_type, "num_pages": len(self.pages)}


def parse_docx(file_path: str) -> ParsedDocument:
    """
    Extracts text from a .docx file.
    Treats each paragraph + table cell as content blocks.
    Groups them into a single logical 'page'.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Expected .docx, got: {path.suffix}")

    doc_obj = Document(str(path))
    blocks: List[str] = []

    # Extract paragraphs
    for para in doc_obj.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    # Extract tables
    for table in doc_obj.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(f"[TABLE ROW] {row_text}")

    doc = ParsedDocument(source=str(path), doc_type="docx")
    doc.pages.append({"page_num": 1, "text": "\n".join(blocks)})

    return doc
